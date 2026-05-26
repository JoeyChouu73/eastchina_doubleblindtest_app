# 飞行员双盲测试评估分析
# Streamlit + Plotly + python-docx

import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

st.set_page_config(page_title="飞行员双盲测试评估分析", layout="wide")

# ------------------ 科目定义（根据Word报告） ------------------
SUBJECTS = {
    '大坡度盘旋': ['高度偏差', '速度偏差', '改出航向', '坡度保持', '滚转速率', '速度保持'],
    '大侧风目视起落': ['起飞抬头率', '一边航迹', 'BANK ANGLE', '三边高度', '三边航迹', '三边宽度', '四边航迹', '四转弯', '下滑线', '五边速度', '稳定意识', 'SINK RATE', '入口高度', '着陆'],
    '非精密进近+中断着陆': ['滑跑方向', '稳定进近', '高距比', '下滑线控制', '中断动作', '中断程序', '入口高度'],
    '中断着陆后发动机失效': ['单发初始姿态', '单发推力', '航迹误差', '坡度控制', '航迹控制', '通信', 'PANPAN'],
    '单发ILS无指引落地': ['1000ft以下五边剖面控制', 'LOC', 'G/S', 'Vapp', '着陆偏出', '剖面控制']
}

# 科目名称映射
SUBJECT_NAMES = {
    '科目1': '大坡度盘旋',
    '科目2': '大侧风目视起落', 
    '科目3': '非精密进近+中断着陆',
    '科目4': '中断着陆后发动机失效',
    '科目5': '单发ILS无指引落地'
}

# 综合考评扣分项
COMPREHENSIVE_ITEMS = ['技术性复飞', '决策失误', 'PULL UP', '盲目蛮干', '冲偏出跑道', '各类音响警戒/警告']


def categorize_deduction_item(col_name):
    """将扣分项归类到具体科目"""
    col_lower = col_name.lower()
    
    for subject, keywords in SUBJECTS.items():
        for keyword in keywords:
            if keyword.lower() in col_lower:
                return subject
    
    # 检查综合考评项
    for item in COMPREHENSIVE_ITEMS:
        if item.lower() in col_lower:
            return '综合考评'
    
    return '其他'


def extract_deductions_with_subjects(df, company_name):
    """提取所有扣分项并归类到科目"""
    df = df.copy()
    
    # 找出所有扣分列
    deduction_cols = []
    for col in df.columns:
        numeric_vals = pd.to_numeric(df[col], errors='coerce')
        if (numeric_vals < 0).any():
            deduction_cols.append(col)
    
    # 计算每条记录的扣分总和
    total_deduction = pd.Series(0, index=df.index)
    deduction_details = {}
    
    # 初始化各科目扣分
    for subject in SUBJECTS.keys():
        df[f'{subject}_扣分'] = 0
    df['综合考评_扣分'] = 0
    df['其他_扣分'] = 0
    
    for col in deduction_cols:
        numeric_vals = pd.to_numeric(df[col], errors='coerce').fillna(0)
        col_deduction = numeric_vals * (numeric_vals < 0)
        total_deduction += col_deduction
        
        # 归类到科目
        subject = categorize_deduction_item(col)
        if subject in SUBJECTS.keys():
            df[f'{subject}_扣分'] += col_deduction
        elif subject == '综合考评':
            df['综合考评_扣分'] += col_deduction
        else:
            df['其他_扣分'] += col_deduction
        
        # 记录每条记录的扣分详情
        for idx in df.index:
            if col_deduction[idx] < 0:
                if idx not in deduction_details:
                    deduction_details[idx] = []
                deduction_details[idx].append({
                    '扣分项': col,
                    '科目': subject,
                    '扣分值': col_deduction[idx]
                })
    
    df['扣分总和'] = total_deduction
    df['最终得分'] = 100 + df['扣分总和']
    df['扣分项数量'] = [len(deduction_details.get(idx, [])) for idx in df.index]
    
    # 计算各科目得分
    for subject in SUBJECTS.keys():
        df[f'{subject}_得分'] = 100 + df[f'{subject}_扣分']
    df['综合考评_得分'] = 100 + df['综合考评_扣分']
    
    return df, deduction_cols


def calculate_subject_stats(df):
    """计算各科目统计数据"""
    subject_stats = {}
    
    for subject in SUBJECTS.keys():
        if f'{subject}_得分' in df.columns:
            subject_stats[subject] = {
                '平均得分': df[f'{subject}_得分'].mean(),
                '平均扣分': -df[f'{subject}_扣分'].mean(),
                '最低分': df[f'{subject}_得分'].min(),
                '最高分': df[f'{subject}_得分'].max(),
                '扣分人数': (df[f'{subject}_扣分'] < 0).sum(),
                '总扣分': -df[f'{subject}_扣分'].sum()
            }
    
    return subject_stats


def analyze_operator_performance(df):
    """按操纵者（机长/副驾驶）分析"""
    if '技术等级' not in df.columns:
        return None
    
    operator_stats = {}
    
    # 识别机长和副驾驶
    captains = df[df['技术等级'].str.contains('机长|教员', na=False)]
    fos = df[df['技术等级'].str.contains('副驾驶', na=False)]
    
    operator_stats['机长'] = {
        '人数': len(captains),
        '平均分': captains['最终得分'].mean() if len(captains) > 0 else 0,
        '最高分': captains['最终得分'].max() if len(captains) > 0 else 0,
        '最低分': captains['最终得分'].min() if len(captains) > 0 else 0
    }
    
    operator_stats['副驾驶'] = {
        '人数': len(fos),
        '平均分': fos['最终得分'].mean() if len(fos) > 0 else 0,
        '最高分': fos['最终得分'].max() if len(fos) > 0 else 0,
        '最低分': fos['最终得分'].min() if len(fos) > 0 else 0
    }
    
    return operator_stats


def analyze_comprehensive_issues(df):
    """分析综合考评问题"""
    comprehensive_data = []
    
    for col in df.columns:
        numeric_vals = pd.to_numeric(df[col], errors='coerce')
        deduction_count = (numeric_vals < 0).sum()
        if deduction_count > 0:
            for item in COMPREHENSIVE_ITEMS:
                if item.lower() in col.lower():
                    comprehensive_data.append({
                        '问题类型': item,
                        '出现次数': deduction_count,
                        '涉及人数': deduction_count
                    })
                    break
    
    if comprehensive_data:
        return pd.DataFrame(comprehensive_data).groupby('问题类型').sum().reset_index().sort_values('出现次数', ascending=False)
    return pd.DataFrame()


def get_top_deductions_by_subject(df, subject, n=3):
    """获取某科目中失分最多的TOP N项"""
    deductions = []
    for col in df.columns:
        numeric_vals = pd.to_numeric(df[col], errors='coerce')
        if (numeric_vals < 0).any():
            subject_match = categorize_deduction_item(col)
            if subject_match == subject:
                deductions.append({
                    '扣分项': col,
                    '出现次数': (numeric_vals < 0).sum(),
                    '总扣分': -numeric_vals[numeric_vals < 0].sum()
                })
    
    return sorted(deductions, key=lambda x: x['出现次数'], reverse=True)[:n]


# ------------------ Word报告（完整版） ------------------
def build_comprehensive_report(all_data, company_stats, subject_stats, operator_stats, 
                               weak_areas_by_subject, comprehensive_issues, top_deductions):
    """生成完整的Word报告，完全覆盖Word报告模板"""
    doc = Document()
    
    # 标题
    title = doc.add_heading('双盲测试数据分析报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"生成时间: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"机型: C909")
    
    # ========== 一、整体情况 ==========
    doc.add_heading('一、整体情况', level=1)
    
    # 1. 参加测试人数
    doc.add_heading('1. 参加测试人数', level=2)
    doc.add_paragraph(f"辖区内C909机型参加本次双盲测试的共有{len(company_stats)}家单位，共计{len(all_data)}人参加了测试。")
    
    # 单位分布表格
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['单位名称', '参加人数', '占比']
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    
    for _, row in company_stats.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row['单位名称'])
        cells[1].text = str(row['人数'])
        cells[2].text = f"{row['人数']/len(all_data)*100:.1f}%"
    
    # 2. 测试分数
    doc.add_heading('2. 测试分数', level=2)
    doc.add_paragraph(f"本次测试满分100分，整体平均分: {all_data['最终得分'].mean():.2f}分")
    doc.add_paragraph(f"最高分: {all_data['最终得分'].max():.0f}分，最低分: {all_data['最终得分'].min():.0f}分")
    
    # 各单位得分表格
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['单位名称', '平均分', '最高分', '最低分', '标准差']
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    
    for _, row in company_stats.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row['单位名称'])
        cells[1].text = f"{row['平均分']:.1f}"
        cells[2].text = f"{row['最高分']:.0f}"
        cells[3].text = f"{row['最低分']:.0f}"
        cells[4].text = f"{row['标准差']:.2f}"
    
    # 机长、副驾驶得分情况
    if operator_stats:
        doc.add_heading('机长、副驾驶得分情况', level=3)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        headers = ['操纵者', '人数', '平均分', '分数范围']
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
        
        for role, stats in operator_stats.items():
            if stats['人数'] > 0:
                cells = table.add_row().cells
                cells[0].text = role
                cells[1].text = str(stats['人数'])
                cells[2].text = f"{stats['平均分']:.1f}"
                cells[3].text = f"{stats['最低分']:.0f}-{stats['最高分']:.0f}"
    
    # ========== 二、科目分析 ==========
    doc.add_heading('二、科目分析', level=1)
    
    # 1. 各航司平均失分统计
    doc.add_heading('1. 各航司平均失分统计', level=2)
    
    for _, row in company_stats.iterrows():
        company = row['单位名称']
        company_data = all_data[all_data['所属单位'] == company]
        company_subject_stats = calculate_subject_stats(company_data)
        
        # 找出失分最多的科目
        max_loss_subject = max(company_subject_stats.items(), key=lambda x: x[1]['平均扣分']) if company_subject_stats else (None, None)
        min_loss_subject = min(company_subject_stats.items(), key=lambda x: x[1]['平均扣分']) if company_subject_stats else (None, None)
        
        doc.add_paragraph(
            f"{company}机型在“{max_loss_subject[0] if max_loss_subject[0] else '未知'}”科目失分最多，"
            f"“{min_loss_subject[0] if min_loss_subject[0] else '未知'}”科目失分最少；"
        )
    
    # 2. 失分TOP 5科目统计
    doc.add_heading('2. 失分TOP 5科目统计', level=2)
    
    all_subject_loss = []
    for subject, stats in subject_stats.items():
        all_subject_loss.append({
            '科目': subject,
            '总失分': stats['总扣分'],
            '平均失分': stats['平均扣分']
        })
    
    top5_subjects = sorted(all_subject_loss, key=lambda x: x['总失分'], reverse=True)[:5]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['科目', '总失分', '平均失分']
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    
    for item in top5_subjects:
        cells = table.add_row().cells
        cells[0].text = item['科目']
        cells[1].text = f"{item['总失分']:.0f}"
        cells[2].text = f"{item['平均失分']:.1f}"
    
    # 3. 各科目失分TOP 3统计
    doc.add_heading('3. 各科目失分TOP 3统计', level=2)
    
    for subject in SUBJECTS.keys():
        top3 = get_top_deductions_by_subject(all_data, subject, 3)
        if top3:
            doc.add_paragraph(f"{subject}科目失分TOP3:", style='List Bullet')
            for item in top3:
                doc.add_paragraph(f"  - {item['扣分项']}: 出现{item['出现次数']}次", style='List Bullet 2')
    
    # 4. 各科目分析
    doc.add_heading('4. 各科目分析', level=2)
    
    for subject in SUBJECTS.keys():
        doc.add_heading(f'（{list(SUBJECTS.keys()).index(subject)+1}）{subject}科目', level=3)
        
        subject_data = all_data.copy()
        subject_col = f'{subject}_得分'
        
        if subject_col in subject_data.columns:
            # I. 整体情况
            doc.add_heading('I. 整体情况', level=4)
            avg_score = subject_data[subject_col].mean()
            doc.add_paragraph(f"该科目平均得分: {avg_score:.1f}分")
            
            # 找出需要加强的关键动作
            top_issues = get_top_deductions_by_subject(all_data, subject, 1)
            if top_issues:
                doc.add_paragraph(f"在“{subject}”科目中“{top_issues[0]['扣分项']}”关键动作需要进一步加强。")
            
            # II. 按操纵者划分
            doc.add_heading('II. 按操纵者划分', level=4)
            if operator_stats:
                for role, stats in operator_stats.items():
                    if stats['人数'] > 0:
                        role_data = all_data[all_data['技术等级'].str.contains(role[:2], na=False)] if role == '机长' else all_data[all_data['技术等级'].str.contains('副驾驶', na=False)]
                        if len(role_data) > 0:
                            role_avg = role_data[subject_col].mean() if subject_col in role_data.columns else 0
                            doc.add_paragraph(f"{role}: 平均分{role_avg:.1f}分，人数{len(role_data)}人")
            
            # III. 各航司得分情况分析
            doc.add_heading('III. 各航司得分情况分析', level=4)
            for _, row in company_stats.iterrows():
                company = row['单位名称']
                company_data = all_data[all_data['所属单位'] == company]
                company_avg = company_data[subject_col].mean() if subject_col in company_data.columns else 0
                
                # 找出该公司在该科目中的主要问题
                top_company_issues = get_top_deductions_by_subject(company_data, subject, 1)
                if top_company_issues:
                    issue_name = top_company_issues[0]['扣分项']
                    doc.add_paragraph(f"{company}在“{subject}”科目中“{issue_name}”关键动作需要加强。")
                else:
                    doc.add_paragraph(f"{company}在“{subject}”科目中表现良好，平均分{company_avg:.1f}分。")
    
    # 5. 综合考评得分分析
    doc.add_heading('5. 综合考评得分分析', level=2)
    
    if not comprehensive_issues.empty:
        doc.add_paragraph("综合考评部分，技术性复飞占比最多，结合前面的科目分析，多数由于前期未建立适合的落地条件所致。")
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        headers = ['问题类型', '出现次数', '涉及人数']
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
        
        for _, row in comprehensive_issues.iterrows():
            cells = table.add_row().cells
            cells[0].text = str(row['问题类型'])
            cells[1].text = str(row['出现次数'])
            cells[2].text = str(row['涉及人数'])
    
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# ------------------ 新增可视化函数 ------------------
def create_company_subject_heatmap(all_data, subjects):
    """各航司科目得分热力图"""
    heatmap_data = []
    for company in all_data['所属单位'].unique():
        company_data = all_data[all_data['所属单位'] == company]
        row = {'单位': company}
        for subject in subjects:
            if f'{subject}_得分' in company_data.columns:
                row[subject] = company_data[f'{subject}_得分'].mean()
        heatmap_data.append(row)
    
    df_heatmap = pd.DataFrame(heatmap_data)
    if len(df_heatmap) > 1 and len(subjects) > 1:
        fig = px.imshow(
            df_heatmap.set_index('单位').values,
            x=subjects,
            y=df_heatmap['单位'].values,
            title="各航司科目得分热力图",
            color_continuous_scale='RdYlGn',
            text_auto='.0f'
        )
        fig.update_layout(font_family="Microsoft YaHei, SimHei", height=400)
        return fig
    return None


def create_operator_comparison_chart(operator_stats):
    """机长/副驾驶对比图"""
    if not operator_stats:
        return None
    
    df = pd.DataFrame([
        {'操纵者': role, '平均分': stats['平均分'], '人数': stats['人数']}
        for role, stats in operator_stats.items() if stats['人数'] > 0
    ])
    
    fig = go.Figure()
    fig.add_trace(go.Bar(
        x=df['操纵者'], y=df['平均分'],
        text=df['平均分'].round(1), textposition='auto',
        marker_color=['#2ecc71', '#3498db']
    ))
    fig.update_layout(
        title="机长与副驾驶得分对比",
        xaxis_title="操纵者", yaxis_title="平均得分",
        font_family="Microsoft YaHei, SimHei", height=400
    )
    return fig


def create_top_deductions_chart(all_data, subject):
    """某科目TOP扣分项图表"""
    top3 = get_top_deductions_by_subject(all_data, subject, 5)
    if not top3:
        return None
    
    df = pd.DataFrame(top3)
    fig = go.Figure()
    fig.add_trace(go.Bar(
        y=df['扣分项'], x=df['出现次数'],
        orientation='h', marker_color='#e74c3c',
        text=df['出现次数'], textposition='outside'
    ))
    fig.update_layout(
        title=f"{subject}科目 - 高频扣分项TOP5",
        xaxis_title="出现次数", yaxis_title="扣分项",
        font_family="Microsoft YaHei, SimHei",
        height=400, margin=dict(l=200)
    )
    return fig


def create_comprehensive_issues_chart(comprehensive_issues):
    """综合考评问题图表"""
    if comprehensive_issues.empty:
        return None
    
    fig = go.Figure()
    fig.add_trace(go.Pie(
        labels=comprehensive_issues['问题类型'],
        values=comprehensive_issues['出现次数'],
        hole=0.4,
        marker_colors=['#e74c3c', '#f39c12', '#e67e22', '#c0392b']
    ))
    fig.update_layout(
        title="综合考评问题分布",
        font_family="Microsoft YaHei, SimHei", height=450
    )
    return fig


# ------------------ 主页面 ------------------
st.title("✈️ 飞行员双盲测试评估分析平台")
st.markdown("基于《华东飞行员双盲测试数据采集表》统一模板，客观扣分计算")

with st.sidebar:
    st.markdown("### 📂 数据上传")
    uploaded_files = st.file_uploader(
        "上传一个或多个Excel文件（每个文件代表一个公司）", 
        type=["xlsx"], 
        accept_multiple_files=True
    )
    
    st.markdown("---")
    st.markdown("### 📌 评分规则")
    st.markdown("""
    - 基础分: **100分**
    - 每出现一次扣分项，扣除相应分数
    - 最终得分 = 100 + 扣分总和（扣分为负值）
    - 分数越高代表表现越好
    """)

if uploaded_files:
    all_data_list = []
    company_names = []
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    for i, file in enumerate(uploaded_files):
        status_text.text(f"正在处理: {file.name}")
        df = load_template(file)
        company_name = file.name.replace('.xlsx', '').replace('_', ' ')
        df, deduction_cols = extract_deductions_with_subjects(df, company_name)
        df['所属单位'] = company_name
        company_names.append(company_name)
        all_data_list.append(df)
        progress_bar.progress((i + 1) / len(uploaded_files))
    
    status_text.text("数据处理完成！")
    all_data = pd.concat(all_data_list, ignore_index=True)
    
    # 计算各类统计
    subject_stats = calculate_subject_stats(all_data)
    operator_stats = analyze_operator_performance(all_data)
    comprehensive_issues = analyze_comprehensive_issues(all_data)
    company_stats = all_data.groupby('所属单位')['最终得分'].agg(
        人数='count', 平均分='mean', 最高分='max', 最低分='min', 标准差='std'
    ).round(2).reset_index()
    company_stats.columns = ['单位名称', '人数', '平均分', '最高分', '最低分', '标准差']
    
    # 筛选器
    companies = ["全部"] + sorted(all_data['所属单位'].unique().tolist())
    tech_levels = ["全部"] + sorted(all_data['技术等级'].dropna().unique().tolist()) if '技术等级' in all_data.columns else ["全部"]
    
    with st.sidebar:
        selected_company = st.selectbox("单位筛选", companies, key="company_filter")
        selected_tech = st.selectbox("技术等级筛选", tech_levels, key="tech_filter")
    
    filtered_data = all_data.copy()
    if selected_company != "全部":
        filtered_data = filtered_data[filtered_data['所属单位'] == selected_company]
    if selected_tech != "全部":
        filtered_data = filtered_data[filtered_data['技术等级'] == selected_tech]
    
    st.success(f"✅ 成功加载 {len(all_data)} 条飞行员记录，来自 {len(company_names)} 个单位")
    
    # 主界面标签页
    tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
        "📊 整体情况", "🏢 单位对比", "📈 科目分析", "🎯 综合考评", "👨‍✈️ 个人详情", "📄 报告输出"
    ])
    
    with tab1:
        st.markdown("### 整体情况")
        
        col1, col2, col3, col4 = st.columns(4)
        col1.metric("参加单位数", len(company_stats))
        col2.metric("参加人数", len(all_data))
        col3.metric("整体平均分", f"{all_data['最终得分'].mean():.1f}")
        col4.metric("最高/最低分", f"{all_data['最终得分'].max():.0f}/{all_data['最终得分'].min():.0f}")
        
        # 机长/副驾驶对比
        if operator_stats:
            st.plotly_chart(create_operator_comparison_chart(operator_stats), use_container_width=True)
        
        # 得分分布
        fig_hist = px.histogram(all_data, x='最终得分', nbins=20, title="整体得分分布",
                                labels={'最终得分': '得分', 'count': '人数'},
                                color_discrete_sequence=['#2ecc71'])
        st.plotly_chart(fig_hist, use_container_width=True)
        
        # 各单位得分箱线图
        fig_box = px.box(all_data, x='所属单位', y='最终得分', title="各单位得分分布对比",
                         points='all', color_discrete_sequence=['#3498db'])
        st.plotly_chart(fig_box, use_container_width=True)
        
        # 统计摘要
        st.markdown("### 统计摘要")
        stats_df = pd.DataFrame({
            "指标": ["平均分", "中位数", "标准差", "最高分", "最低分", "合格率(≥80)", "优秀率(≥90)"],
            "数值": [
                f"{all_data['最终得分'].mean():.2f}",
                f"{all_data['最终得分'].median():.2f}",
                f"{all_data['最终得分'].std():.2f}",
                f"{all_data['最终得分'].max():.0f}",
                f"{all_data['最终得分'].min():.0f}",
                f"{(all_data['最终得分'] >= 80).mean()*100:.1f}%",
                f"{(all_data['最终得分'] >= 90).mean()*100:.1f}%"
            ]
        })
        st.dataframe(stats_df, use_container_width=True)
    
    with tab2:
        st.markdown("### 各单位对比分析")
        st.dataframe(company_stats, use_container_width=True, hide_index=True)
        
        col1, col2 = st.columns(2)
        with col1:
            fig_bar = px.bar(company_stats, x='单位名称', y='平均分', title="各单位平均分对比",
                            text='平均分', color_discrete_sequence=['#2ecc71'])
            fig_bar.update_traces(textposition='outside')
            st.plotly_chart(fig_bar, use_container_width=True)
        
        with col2:
            fig_radar = create_company_radar(company_stats)
            if fig_radar:
                st.plotly_chart(fig_radar, use_container_width=True)
        
        # 各航司科目得分热力图
        fig_heatmap = create_company_subject_heatmap(all_data, list(SUBJECTS.keys()))
        if fig_heatmap:
            st.plotly_chart(fig_heatmap, use_container_width=True)
        
        # 各单位详细分析
        st.markdown("### 各单位详细分析")
        for company in company_stats['单位名称']:
            with st.expander(f"📊 {company} 详细分析"):
                company_data = all_data[all_data['所属单位'] == company]
                company_subject_stats = calculate_subject_stats(company_data)
                
                # 找出失分最多和最少的科目
                if company_subject_stats:
                    max_loss = max(company_subject_stats.items(), key=lambda x: x[1]['平均扣分'])
                    min_loss = min(company_subject_stats.items(), key=lambda x: x[1]['平均扣分'])
                    st.info(f"失分最多: {max_loss[0]} (平均扣分{max_loss[1]['平均扣分']:.1f}) | 失分最少: {min_loss[0]} (平均扣分{min_loss[1]['平均扣分']:.1f})")
                
                # 该公司各科目得分
                subject_scores = []
                for subject in SUBJECTS.keys():
                    if f'{subject}_得分' in company_data.columns:
                        subject_scores.append({
                            '科目': subject,
                            '平均分': company_data[f'{subject}_得分'].mean(),
                            '扣分人数': (company_data[f'{subject}_扣分'] < 0).sum()
                        })
                if subject_scores:
                    st.dataframe(pd.DataFrame(subject_scores), use_container_width=True, hide_index=True)
    
    with tab3:
        st.markdown("### 科目分析")
        
        # 各科目平均分雷达图
        subject_radar_data = {k: v['平均得分'] for k, v in subject_stats.items()}
        if subject_radar_data:
            fig_radar_subject = go.Figure()
            fig_radar_subject.add_trace(go.Scatterpolar(
                r=list(subject_radar_data.values()),
                theta=list(subject_radar_data.keys()),
                fill='toself', name='平均得分', line_color='#2ecc71'
            ))
            fig_radar_subject.update_layout(
                polar=dict(radialaxis=dict(visible=True, range=[0, 100])),
                title="各科目平均得分雷达图", height=450
            )
            st.plotly_chart(fig_radar_subject, use_container_width=True)
        
        # 失分TOP5科目
        st.markdown("#### 失分TOP5科目")
        top5_subjects = sorted(subject_stats.items(), key=lambda x: x[1]['总扣分'], reverse=True)[:5]
        top5_df = pd.DataFrame([{
            '科目': s[0], '总失分': s[1]['总扣分'], '平均失分': s[1]['平均扣分']
        } for s in top5_subjects])
        st.dataframe(top5_df, use_container_width=True, hide_index=True)
        
        # 各科目详细分析
        st.markdown("#### 各科目详细分析")
        for subject in SUBJECTS.keys():
            with st.expander(f"📌 {subject}科目分析"):
                if subject in subject_stats:
                    stats = subject_stats[subject]
                    col1, col2, col3 = st.columns(3)
                    col1.metric("平均得分", f"{stats['平均得分']:.1f}")
                    col2.metric("平均扣分", f"{stats['平均扣分']:.1f}")
                    col3.metric("扣分人数", stats['扣分人数'])
                    
                    # 该科目TOP扣分项
                    fig_top = create_top_deductions_chart(all_data, subject)
                    if fig_top:
                        st.plotly_chart(fig_top, use_container_width=True)
                    
                    # 各航司在该科目的表现
                    st.markdown("**各航司表现**")
                    company_subject = []
                    for company in all_data['所属单位'].unique():
                        company_data = all_data[all_data['所属单位'] == company]
                        if f'{subject}_得分' in company_data.columns:
                            company_subject.append({
                                '单位': company,
                                '平均分': company_data[f'{subject}_得分'].mean(),
                                '扣分人数': (company_data[f'{subject}_扣分'] < 0).sum()
                            })
                    if company_subject:
                        st.dataframe(pd.DataFrame(company_subject), use_container_width=True, hide_index=True)
    
    with tab4:
        st.markdown("### 综合考评分析")
        
        if not comprehensive_issues.empty:
            st.plotly_chart(create_comprehensive_issues_chart(comprehensive_issues), use_container_width=True)
            st.dataframe(comprehensive_issues, use_container_width=True, hide_index=True)
            st.info("💡 综合考评部分，技术性复飞占比最多，结合前面的科目分析，多数由于前期未建立适合的落地条件所致。")
        else:
            st.info("未发现综合考评扣分项")
    
    with tab5:
        st.markdown("### 个人详情查询")
        
        col1, col2 = st.columns(2)
        with col1:
            company_filter = st.selectbox("选择单位", ["全部"] + sorted(all_data['所属单位'].unique().tolist()), key="personal_company")
        with col2:
            if '技术等级' in all_data.columns:
                tech_filter = st.selectbox("选择技术等级", ["全部"] + sorted(all_data['技术等级'].dropna().unique().tolist()), key="personal_tech")
            else:
                tech_filter = "全部"
        
        search_name = st.text_input("🔍 搜索姓名", "")
        
        personal_data = all_data.copy()
        if company_filter != "全部":
            personal_data = personal_data[personal_data['所属单位'] == company_filter]
        if tech_filter != "全部":
            personal_data = personal_data[personal_data['技术等级'] == tech_filter]
        if search_name:
            personal_data = personal_data[personal_data['姓名'].astype(str).str.contains(search_name, na=False)]
        
        # 显示表格
        display_cols = ['姓名', '所属单位', '技术等级', '最终得分', '扣分项数量'] + [f'{s}_得分' for s in SUBJECTS.keys() if f'{s}_得分' in personal_data.columns]
        display_cols = [c for c in display_cols if c in personal_data.columns]
        
        st.dataframe(personal_data[display_cols].sort_values('最终得分', ascending=False), 
                     use_container_width=True, hide_index=True)
        
        # 个人详细分析
        if len(personal_data) > 0:
            st.markdown("### 📋 个人详细分析")
            selected_pilot = st.selectbox("选择飞行员", personal_data['姓名'].tolist())
            pilot_data = personal_data[personal_data['姓名'] == selected_pilot].iloc[0]
            
            col1, col2, col3, col4 = st.columns(4)
            col1.metric("最终得分", f"{pilot_data['最终得分']:.1f}")
            col2.metric("扣分项数量", pilot_data['扣分项数量'])
            col3.metric("所属单位", pilot_data['所属单位'])
            if '技术等级' in pilot_data.index:
                col4.metric("技术等级", pilot_data['技术等级'])
            
            # 各科目得分
            st.markdown("#### 各科目得分")
            subject_scores = []
            for subject in SUBJECTS.keys():
                if f'{subject}_得分' in pilot_data.index:
                    subject_scores.append({'科目': subject, '得分': pilot_data[f'{subject}_得分']})
            if subject_scores:
                st.dataframe(pd.DataFrame(subject_scores), use_container_width=True, hide_index=True)
            
            # 扣分详情
            st.markdown("#### 扣分详情")
            pilot_deductions = []
            for col in all_data.columns:
                if col not in ['姓名', '所属单位', '技术等级', '最终得分', '扣分项数量', '扣分总和']:
                    val = pd.to_numeric(pilot_data.get(col, 0), errors='coerce')
                    if val < 0:
                        pilot_deductions.append({'扣分项': col, '扣分值': val})
            
            if pilot_deductions:
                st.dataframe(pd.DataFrame(pilot_deductions), use_container_width=True, hide_index=True)
            else:
                st.success("该飞行员无扣分项，表现完美！")
    
    with tab6:
        st.markdown("### 📄 生成分析报告")
        
        # 生成完整报告
        report_bio = build_comprehensive_report(
            all_data, company_stats, subject_stats, operator_stats,
            None, comprehensive_issues, None
        )
        
        st.download_button(
            "📥 下载完整Word报告",
            report_bio,
            f"双盲测试数据分析报告_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        # CSV数据导出
        csv_data = all_data.to_csv(index=False, encoding='utf-8-sig')
        st.download_button(
            "📥 下载CSV数据",
            csv_data,
            f"飞行员评估数据_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.csv",
            mime="text/csv"
        )
        
        # 报告预览
        st.markdown("### 报告预览")
        preview_text = f"""
        ═══════════════════════════════════════════════════════════
                       双盲测试数据分析报告
        ═══════════════════════════════════════════════════════════
        
        一、整体情况
        
        1. 参加测试人数
        辖区内C909机型参加本次双盲测试的共有{len(company_stats)}家单位，共计{len(all_data)}人参加了测试。
        
        2. 测试分数
        本次测试满分100分，整体平均分: {all_data['最终得分'].mean():.2f}分
        最高分: {all_data['最终得分'].max():.0f}分，最低分: {all_data['最终得分'].min():.0f}分
        
        机长平均分: {operator_stats.get('机长', {}).get('平均分', 0):.1f}分
        副驾驶平均分: {operator_stats.get('副驾驶', {}).get('平均分', 0):.1f}分
        
        二、科目分析
        
        1. 各航司平均失分统计
        """
        for _, row in company_stats.iterrows():
            preview_text += f"\n  {row['单位名称']}: 平均分{row['平均分']:.1f}分"
        
        st.text_area("报告预览", preview_text, height=400)

else:
    st.info("👈 请先从左侧上传一个或多个Excel文件（每个文件代表一个公司）")
    
    with st.expander("📖 文件格式说明"):
        st.markdown("""
        **支持的文件格式**：
        - 按《华东飞行员双盲测试数据采集表》模板填写的Excel文件
        - 每个文件代表一个公司的数据
        
        **系统会自动**：
        1. 按科目（大坡度盘旋、大侧风目视起落、非精密进近+中断着陆、中断着陆后发动机失效、单发ILS无指引落地）分类统计
        2. 分析各航司、机长/副驾驶得分情况
        3. 识别高频扣分项，定位薄弱环节
        4. 生成完整的Word分析报告
        """)

st.markdown("---")
st.caption("飞行员双盲测试评估分析 | 基于华东局B737飞行技能评估标准 | 完全覆盖Word报告分析维度")